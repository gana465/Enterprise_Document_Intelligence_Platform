"""
utils.py
----------------------------------------
Utility Functions
Enterprise Document Intelligence Platform
----------------------------------------
"""

import hashlib
import mimetypes
import os
import re
import shutil
from pathlib import Path

import fitz
import docx

from config import (
    SUPPORTED_EXTENSIONS,
    MAX_UPLOAD_SIZE_MB,
    UPLOAD_FOLDER
)

# --------------------------------------------------------
# File Validation
# --------------------------------------------------------

def allowed_file(filename: str) -> bool:

    extension = Path(filename).suffix.lower()

    return extension in SUPPORTED_EXTENSIONS


# --------------------------------------------------------
# MIME Type
# --------------------------------------------------------

def get_mime_type(file_path):

    mime, _ = mimetypes.guess_type(file_path)

    return mime or "application/octet-stream"


# --------------------------------------------------------
# File Size
# --------------------------------------------------------

def file_size_mb(file_path):

    size = os.path.getsize(file_path)

    return round(size / (1024 * 1024), 2)


# --------------------------------------------------------
# Check Upload Size
# --------------------------------------------------------

def validate_size(file_path):

    return file_size_mb(file_path) <= MAX_UPLOAD_SIZE_MB


# --------------------------------------------------------
# SHA256 Hash
# --------------------------------------------------------

def sha256(file_path):

    h = hashlib.sha256()

    with open(file_path, "rb") as f:

        while True:

            chunk = f.read(8192)

            if not chunk:
                break

            h.update(chunk)

    return h.hexdigest()


# --------------------------------------------------------
# Save Uploaded File
# --------------------------------------------------------

def save_uploaded_file(uploaded_file):

    destination = UPLOAD_FOLDER / uploaded_file.name

    with open(destination, "wb") as f:

        shutil.copyfileobj(uploaded_file, f)

    return destination


# --------------------------------------------------------
# Safe Filename
# --------------------------------------------------------

def sanitize_filename(filename):

    filename = filename.replace(" ", "_")

    filename = re.sub(
        r"[^A-Za-z0-9._-]",
        "",
        filename
    )

    return filename


# --------------------------------------------------------
# PDF Extraction
# --------------------------------------------------------

def extract_pdf(path):

    doc = fitz.open(path)

    text = ""

    for page in doc:

        text += page.get_text()

    doc.close()

    return text


# --------------------------------------------------------
# DOCX Extraction
# --------------------------------------------------------

def extract_docx(path):

    document = docx.Document(path)

    paragraphs = []

    for p in document.paragraphs:

        paragraphs.append(p.text)

    return "\n".join(paragraphs)


# --------------------------------------------------------
# TXT Extraction
# --------------------------------------------------------

def extract_txt(path):

    with open(
        path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as f:

        return f.read()


# --------------------------------------------------------
# Automatic Extraction
# --------------------------------------------------------

def extract_text(path):

    extension = Path(path).suffix.lower()

    if extension == ".pdf":

        return extract_pdf(path)

    if extension == ".docx":

        return extract_docx(path)

    if extension == ".txt":

        return extract_txt(path)

    return ""


# --------------------------------------------------------
# Clean Text
# --------------------------------------------------------

def clean_text(text):

    text = re.sub(r"\s+", " ", text)

    text = text.strip()

    return text


# --------------------------------------------------------
# Word Count
# --------------------------------------------------------

def word_count(text):

    return len(text.split())


# --------------------------------------------------------
# Character Count
# --------------------------------------------------------

def character_count(text):

    return len(text)


# --------------------------------------------------------
# Page Count
# --------------------------------------------------------

def page_count(file_path):

    from pathlib import Path
import fitz
from docx import Document


def page_count(file_path):

    # Convert Path object to string if needed
    file_path = Path(file_path)

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":

        try:
            pdf = fitz.open(file_path)
            pages = len(pdf)
            pdf.close()
            return pages

        except Exception:
            return 0

    elif suffix == ".docx":

        try:
            doc = Document(file_path)

            # Approximate page count
            return max(1, len(doc.paragraphs) // 40)

        except Exception:
            return 1

    else:
        return 1


# --------------------------------------------------------
# Preview
# --------------------------------------------------------

def preview(text, words=75):

    tokens = text.split()

    if len(tokens) <= words:

        return text

    return " ".join(tokens[:words]) + " ..."


# --------------------------------------------------------
# Metadata
# --------------------------------------------------------

def metadata(file_path):

    from pathlib import Path


def metadata(file_path):

    file_path = Path(file_path)

    return {

        "filename": file_path.name,

        "size": round(
            file_path.stat().st_size / (1024 * 1024),
            2
        ),

        "pages": page_count(file_path)

    }